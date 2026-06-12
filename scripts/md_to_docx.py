#!/usr/bin/env python3
"""Convert output/sop_sharepoint_section.md to a Word .docx.

Maps Markdown to Word styles as the SOP section itself specifies:
  `#`  -> Heading 1
  `##` -> Heading 2
Bullets (with 2-space nesting), inline **bold** / *italic*, the leading
blockquote note, and `---` rules are handled. Stdlib + python-docx only.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

SRC = Path("output/sop_sharepoint_section.md")
OUT = Path("output/sop_sharepoint_section.docx")

# Split a line of markdown into (text, bold, italic) runs.
_TOKEN = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")


def add_runs(paragraph, text):
    pos = 0
    for m in _TOKEN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            paragraph.add_run(m.group(1)).bold = True
        else:
            paragraph.add_run(m.group(2)).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def main():
    if not SRC.exists():
        sys.exit(f"missing source: {SRC}")
    raw_lines = SRC.read_text(encoding="utf-8").splitlines()

    # The source is hard-wrapped: join continuation lines back into one
    # logical line so paragraphs and bullets aren't split (and so inline
    # **bold**/*italic* spanning a wrap still parses).
    lines = []

    def is_continuation(s):
        # A wrapped line of body text: not blank, rule, heading, quote, or a
        # new bullet marker. Such a line folds back into the previous one.
        return bool(
            s.strip()
            and s.strip() != "---"
            and not s.startswith("#")
            and not s.startswith(">")
            and not re.match(r"^\s*-\s+", s)
        )

    def can_continue(s):
        # Previous logical line a continuation may attach to: a paragraph or a
        # bullet item — anything but blank, rule, or heading.
        return bool(s.strip() and s.strip() != "---" and not s.startswith("#"))

    for raw in raw_lines:
        if lines and is_continuation(raw) and can_continue(lines[-1]):
            lines[-1] = lines[-1].rstrip() + " " + raw.strip()
        else:
            lines.append(raw)

    doc = Document()

    quote_buf = []  # accumulate leading blockquote lines

    def flush_quote():
        if not quote_buf:
            return
        p = doc.add_paragraph(style="Intense Quote")
        add_runs(p, " ".join(quote_buf))
        quote_buf.clear()

    for raw in lines:
        line = raw.rstrip()

        if line.startswith(">"):
            quote_buf.append(line.lstrip("> ").rstrip())
            continue
        flush_quote()

        if not line.strip():
            continue
        if line.strip() == "---":
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^\s*-\s+", line):
            indent = len(line) - len(line.lstrip(" "))
            level = indent // 2  # 0 -> top, 2-space -> nested
            style = "List Bullet" if level == 0 else f"List Bullet {min(level + 1, 3)}"
            text = re.sub(r"^\s*-\s+", "", line)
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            add_runs(p, text)
        else:
            p = doc.add_paragraph()
            add_runs(p, line.strip())

    flush_quote()
    doc.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
